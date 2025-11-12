# streamlit_incident_builder.py

import streamlit as st
import requests
import json
import base64
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================
# CONFIG DEFAULTS (you can override in the Streamlit sidebar)
# ============================================================
DEFAULT_GROK_MODEL = "grok-4-fast-reasoning"  # adjust to your deployed model


# ============================================================
# MOCK INCIDENTS (so you don't have to type everything)
# ============================================================

def get_mock_user_data_case1():
    """
    Mock case similar to your Kelt / partial bump / continued inflow scenario.
    """
    return {
        "cir_number": "CIR-25-99",
        "date_of_report": "2025-04-01",
        "revision": "0",
        "customer": "Test Exploration Ltd.",
        "rig": "Test Rig 101",
        "surface_location": "01-01-01-01W5",
        "uwi": "100/01-01-001-01W5",
        "author": "Test Engineer, P.Eng",
        "title_line": "Mock Citadel No Bump & Pressure Test",

        "string_desc": '4-1/2" 21.0kg/m L80 LTC x 5-1/2" 32.0kg/m L80 LTC production casing',
        "accessories": [
            "15K Citadel SV Float Shoe",
            "15K Citadel SV Float Collar",
            "Citadel Latch Cement Plug",
            "2 x NCS Toe Ports",
            "60 x NCS Frac Sleeves",
            "NCS Airlock"
        ],

        # Pre-job & well geometry
        "hole_size_mm": 171,
        "td_mmd": 5300.0,
        "set_depth_mmd": 5297.5,
        "pre_cement_notes": (
            "No sustained inflow prior to cement. Standpipe pressure became erratic "
            "during cleanup while circulating across the airlock."
        ),
        "circulated_volume_m3": 120.0,

        # Cement job
        "cement_lead_m3": 80.0,
        "cement_tail_m3": None,
        "displacement_pumped_m3": 46.0,
        "pump_rate_m3_per_min": 1.0,

        # Pressures / flowback
        "fcp_mpa": 19.0,
        "bump_pressure_mpa": 23.0,
        "bledoff_to_mpa": 14.0,
        "flowback_volume": "≈1.1 m³ and continued to flow at low rate",

        # Volume / depth summary
        "volume_table": {
            "well_td_m": 5300.0,
            "well_tvd_m": 1650.0,
            "shoe_depth_m": 5297.0,
            "float_collar_depth_m": 5284.0,
            "airlock_depth_m": 2600.0,
            "crossover_depth_m": 1700.0,
            "casing_vol_nominal_m3": 45.5,
            "casing_vol_min_m3": 43.7,
            "casing_vol_max_m3": 47.3,
            "volume_to_airlock_m3": 25.3,
            "buoyant_volume_nom_m3": 20.3,
            "displacement_pumped_m3": 46.0,
            "excess_to_surface_m3": None,
        },

        # Post-job information
        "post_job": {
            "retest_pressure_mpa": "18–20",
            "bridge_plug_hold_mpa": "21",
            "bridge_plug_hold_min": "10",
            "tag_depth_m": "N/A",
            "squeeze_summary": "No squeeze performed on this mock case"
        },

        # Optional mismatch info
        "mismatch_receptacle_receptacle_id_in": "",
        "mismatch_receptacle_plug_nose_id_in": "",
    }


def get_mock_user_data_case2():
    """
    Mock incident representing a failure likely due to debris /
    mismatched hardware / damaged float equipment before cementing.
    """
    return {
        "cir_number": "CIR-25-100",
        "date_of_report": "2025-04-05",
        "revision": "0",
        "customer": "Frontier Energy Ltd.",
        "rig": "Precision 555",
        "surface_location": "14-32-79-10W6",
        "uwi": "100/14-32-079-10W6",
        "author": "Field Engineer, P.Eng",
        "title_line": "Plug Did Not Bump – Suspected Debris / Collar Restriction",

        "string_desc": '5-1/2" 26.8kg/m P110 LTC production casing',
        "accessories": [
            "15K Citadel Float Shoe",
            "15K Citadel Float Collar",
            "Citadel Latch Plug",
            "Buoyancy Sub",
            "Toe Sleeve Assembly",
            "Airlock CBS"
        ],

        # Pre-job & well geometry
        "hole_size_mm": 222,
        "td_mmd": 4560.0,
        "set_depth_mmd": 4557.5,
        "pre_cement_notes": (
            "During circulation, SPP became erratic and surged several times as fragments "
            "of glass and composite debris were circulated out from the buoyancy sub. "
            "Returns were intermittently gas-cut prior to cementing."
        ),
        "circulated_volume_m3": 95.0,

        # Cement job
        "cement_lead_m3": 65.0,
        "cement_tail_m3": 15.0,
        "displacement_pumped_m3": 37.5,
        "pump_rate_m3_per_min": 0.8,

        # Pressures / flowback
        "fcp_mpa": 17.2,
        "bump_pressure_mpa": 25.8,
        "bledoff_to_mpa": 0.0,
        "flowback_volume": "minimal flowback (<0.1 m³), pressure dropped immediately after bump attempt",

        # Volume / depth summary
        "volume_table": {
            "well_td_m": 4560.0,
            "well_tvd_m": 1485.0,
            "shoe_depth_m": 4558.0,
            "float_collar_depth_m": 4543.0,
            "airlock_depth_m": 2680.0,
            "crossover_depth_m": 1670.0,
            "casing_vol_nominal_m3": 36.8,
            "casing_vol_min_m3": 35.4,
            "casing_vol_max_m3": 38.2,
            "volume_to_airlock_m3": 20.2,
            "buoyant_volume_nom_m3": 16.4,
            "displacement_pumped_m3": 37.5,
            "excess_to_surface_m3": None,
        },

        # Post-job info
        "post_job": {
            "retest_pressure_mpa": "Could not pressure test; zero isolation",
            "bridge_plug_hold_mpa": "N/A",
            "bridge_plug_hold_min": "N/A",
            "tag_depth_m": "N/A",
            "squeeze_summary": "Planned remedial cement squeeze to restore isolation across shoe track"
        },

        # Mismatch suspected
        "mismatch_receptacle_receptacle_id_in": "4.000",
        "mismatch_receptacle_plug_nose_id_in": "3.875",
    }


# ============================================================
# ROOT-CAUSE TEMPLATE PARAGRAPHS (your approved wording)
# ============================================================

def cause_incorrect_pumping_volume(data):
    vt = data["volume_table"]
    nominal = vt.get("casing_vol_nominal_m3", "N/A")
    disp = data.get("displacement_pumped_m3", "N/A")
    fcp = data.get("fcp_mpa", "N/A")
    bump = data.get("bump_pressure_mpa", "N/A")
    bled = data.get("bledoff_to_mpa", "N/A")

    return f"""Incorrect Pumping Volumes & Leak Above Float Collar

Review of drilling reports, cement charts, and Pason EDR data shows that the correct theoretical displacement (nominal) volume was pumped to bump the plug. The plug was reported to partially bump at {disp}m³ (nominal {nominal}m³); however pressure did not hold and bled off. FCP was {fcp}MPa, bump {bump}MPa, then bled down to {bled}MPa.

API 5CT allows tolerance on pipe wall thickness, which can change actual casing volume versus nominal min/max. If you pump only the nominal volume with no allowance for tolerance, aeration, and compressibility, the plug might not fully land and latch.

Given the Pason / cement data, it appears the correct volume was displaced and a partial bump was observed. We do not believe total displacement volume is the cause of the incident.
""".strip()


def cause_compressibility_ballooning(data, compressibility_outcome: str):
    disp = data.get("displacement_pumped_m3", "N/A")
    bump = data.get("bump_pressure_mpa", "N/A")
    flowback = data.get("flowback_volume", "N/A")

    calc_theoretical_L = "200"
    calc_thermal_L = "50"

    if compressibility_outcome == "plausible":
        conclusion = (
            "These values are consistent with the observed flowback "
            f"({flowback}). Based on this, floats appear to have held, "
            "and the observed bleedoff/flowback can be explained by "
            "normal compressibility and warm-back effects."
        )
    else:
        conclusion = (
            f"However, approximately {flowback} came back and continued. "
            "That volume exceeds what we consider normal compressibility / "
            "aeration / thermal effects. That means the negative inflow "
            "test actually failed: there is communication in the string, "
            "and integrity above the float collar is in question."
        )

    return f"""Fluid Compressibility, Casing Ballooning, and Thermal Expansion

We calculated compressibility for {disp}m³ at ~{bump}MPa surface-applied pressure. The theoretical compressed volume is approximately {calc_theoretical_L} L, plus an additional {calc_thermal_L} L from conservative thermal expansion.

{conclusion}
""".strip()


def cause_failure_prior_to_cementing(data):
    cleanup_vol = data.get("circulated_volume_m3", "N/A")

    return f"""Failure Prior to Cementing

Casing was run to TD and the buoyancy / airlock sub was burst. Circulation to clean up removed debris, but standpipe pressure became erratic as large glass/ceramic fragments met the plug receptacle and float equipment. We have repeatedly seen cases where buoyancy discs do not fully rupture into fine particles. Instead, large shards or “cored out” rings act like a hard seat that can interfere with plug travel, damage seal faces, or bridge at the crossover.

Based on the Pason EDR data and cleanup volumes (~{cleanup_vol}m³ circulated prior to cementing), we believe debris from the buoyancy sub likely restricted or damaged the plug landing area prior to cementing. This can prevent a proper latch and can explain later high pump pressure spikes and abnormal bump behavior.
""".strip()


def cause_mismatched_receptacle(user_data):
    rec_id = user_data.get("mismatch_receptacle_receptacle_id_in") or "N/A"
    plug_id = user_data.get("mismatch_receptacle_plug_nose_id_in") or "N/A"

    return f"""Incorrect Nose Assembly / Mismatched Plug & Collar

Investigation into orders, delivery tickets, and yard pulls shows the incorrect plugs were shipped to location. The float collar receptacle ID was {rec_id}", while the shipped plug nose was designed for {plug_id}". 

If the larger-ID receptacle is run with the smaller-ID plug nose (or vice versa), the plug can physically land but it will not seal. You may see a “bump” indication, but you will not get a sustained positive pressure test.

We have already flagged this QA/QC issue between yards and implemented engineering, design, and QAQC policy changes so mismatched nose/receptacle sets do not get combined on location again.
""".strip()


def cause_debris_on_collar(data):
    bump = data.get("bump_pressure_mpa", "N/A")
    bled = data.get("bledoff_to_mpa", "N/A")
    retest = data.get("post_job", {}).get("retest_pressure_mpa", "N/A")
    hold_press = data.get("post_job", {}).get("bridge_plug_hold_mpa", "N/A")
    hold_min = data.get("post_job", {}).get("bridge_plug_hold_min", "N/A")

    return f"""Debris Below Plug

During displacement we saw a clear bump at ~{bump}MPa but the pressure would not hold above ~{bled}MPa. The pressure then bled down and stabilized instead of dropping straight to zero, which suggests bypass across the plug rather than an open surface valve.

Afterward, we attempted to pressure test casing using rig pumps and again saw bleedoff near {retest}MPa. We then set a bridge plug above the float collar and successfully tested that upper section of casing to {hold_press}MPa for {hold_min} minutes. That proves casing above the collar was sound.

This behavior is consistent with debris sitting on the float collar landing face before the plug arrived. That debris prevents the plug from sitting flush, lets displacement fluid bypass, and can essentially evacuate cement from the shoe track. We have observed this failure mode in past jobs.
""".strip()


def cause_third_party_integrity(data):
    bled = data.get("bledoff_to_mpa", "N/A")
    flowback = data.get("flowback_volume", "N/A")

    return f"""Failure of Third-Party Casing Accessories & Connections

We landed on calculated displacement volume and saw an apparent bump, but we could not hold a positive pressure test. The string repeatedly bled down and stabilized around {bled}MPa. In addition, post-job data indicates inflow that exceeded normal compressibility / thermal expansion (reported ~{flowback}).

When the plug lands on volume but we cannot maintain pressure, and we later observe inflow, the most probable explanation is loss of integrity somewhere above the float collar: a casing connection, frac/toe port, or other third-party accessory. The pattern here is not consistent with float equipment failure.
""".strip()


ROOT_CAUSE_BLOCK_BUILDERS = {
    "incorrect_pumping_volume": lambda d: cause_incorrect_pumping_volume(d),
    "compressibility_ballooning": lambda d: cause_compressibility_ballooning(
        d, d.get("compressibility_outcome", "exceeds_normal")
    ),
    "failure_prior_to_cementing": lambda d: cause_failure_prior_to_cementing(d),
    "mismatched_receptacle": lambda d: cause_mismatched_receptacle(d),
    "debris_on_collar": lambda d: cause_debris_on_collar(d),
    "third_party_integrity": lambda d: cause_third_party_integrity(d),
}


# ============================================================
# Grok call – with optional images
# ============================================================

def encode_uploaded_images(uploaded_files):
    """
    Turn Streamlit uploaded files into a list of dicts:
    [
      {
        "filename": ...,
        "mime_type": ...,
        "b64": "<base64 string>"
      }
    ]
    """
    images = []
    for f in uploaded_files:
        raw = f.read()
        if not raw:
            continue
        b64 = base64.b64encode(raw).decode("utf-8")
        mime = f.type or "image/png"
        images.append({
            "filename": f.name,
            "mime_type": mime,
            "b64": b64,
        })
    return images


def generate_ai_full_report(user_data: dict, api_key: str, model: str, images=None) -> dict:
    """
    Ask Grok to:
    - pick applicable root cause modules from our allowed list
    - classify compressibility outcome
    - write LONG, detailed narrative sections
    - explain reasoning in overall_cause_analysis
    - optionally consider uploaded images (EDR screenshots, etc.)

    images: list of {"filename", "mime_type", "b64"}
    """

    if images is None:
        images = []

    # allowed_modules = [
    #     "incorrect_pumping_volume",
    #     "compressibility_ballooning",
    #     "failure_prior_to_cementing",
    #     "mismatched_receptacle",
    #     "debris_on_collar",
    #     "third_party_integrity"
    # ]

    allowed_modules = [
    # Core modules
    "incorrect_pumping_volume",
    "compressibility_ballooning",
    "failure_prior_to_cementing",
    "mismatched_receptacle",
    "debris_on_collar",
    "third_party_integrity",

    # From CIR-25-19+ to 31 series
    "plug_bumped_early_not_latched",
    "valve_or_plunger_sticking",
    "plug_damaged_or_deformed",
    "float_valve_leakback",
    "third_party_port_leakage",
    "plug_misalignment_in_latch",
    "premature_float_activation",
    "hydraulic_lock_during_displacement",
    "damaged_float_shoe_face",
    "gas_cut_fluid_in_cement",
    "operational_interruption_during_bump",
    "mechanical_damage_during_drillout",

    # From CIR-25-12 to 18 series
    "no_plug_bump_no_test",
    "plug_landed_but_not_tested",
    "float_valve_did_not_close",
    "float_valve_stuck_open",
    "plug_receptacle_tolerance_issue",
    "underdisplacement_due_to_rate_cutback",
    "cement_channel_or_bypass",
    "fractured_float_body",
    "plug_assembly_contamination",
    "flowback_through_toe",
    "premature_cement_cutoff",
    "no_wits_data_verification",
    "cb_sub_failed_burst_disk",

    "float_shoe_damaged_during_run",
    "float_valve_plugged_by_debris",
    "float_valve_not_fully_closing",
    "cold_weather_surface_issue",
    "underdisplacement_due_to_operational_interruptions",
    "plug_bypassed_or_not_latched",
    "float_collar_or_shoe_leakage",
    "buoyancy_disc_failure",

    "no_plug_bump_inflow",
    "plug_bypassed_or_damaged",
    "valve_or_plunger_damaged_during_run",
    "float_valve_leak_post_cement",
    "staged_pump_shutdown_pressure_spike",
    "segmented_plug_design_limitation",
    "hydraulic_shock_during_burst",
    "third_party_integrity_failure",
    "compressibility_exceeds_expected",
    "pressure_fluctuation_during_staging",
    "plug_seat_misalignment",

    "incomplete_air_displacement",
    "compressibility_and_thermal_expansion",
    "burst_disc_debris_bridging_valves",
    "third_party_tool_communication",
    "debris_from_casing_or_lcm",
    "plug_not_seated_in_float_collar",
    "connection_leak_above_float",
    "incorrect_pumping_tolerance_margin",
    "burst_port_or_toe_sleeve_misconfiguration",
    "float_valve_partial_closure",
]


    # Flatten relevant facts into text bullets
    facts_lines = []
    for key, val in user_data.items():
        if key in ["volume_table", "post_job"]:
            continue
        facts_lines.append(f"- {key}: {val}")
    for k, v in user_data["volume_table"].items():
        facts_lines.append(f"- volume_table.{k}: {v}")
    for k, v in user_data["post_job"].items():
        facts_lines.append(f"- post_job.{k}: {v}")
    facts_lines.append(
        f"- mismatch_receptacle_receptacle_id_in: {user_data.get('mismatch_receptacle_receptacle_id_in')}"
    )
    facts_lines.append(
        f"- mismatch_receptacle_plug_nose_id_in: {user_data.get('mismatch_receptacle_plug_nose_id_in')}"
    )

    facts_blob = "\n".join(facts_lines)

    # Build system instructions
    system_instruction = f"""
You are an experienced completions/cementing engineer writing an internal incident investigation report for CCAI.

Tone rules:
- Factual, past-tense, neutral, technical.
- Use MPa for pressures, m³ for volumes, mMD for measured depth.
- Do not assign personal blame. Use language like "suggests", "indicates", "appears".
- Audience is engineering / production / management.

You MUST reason from ONLY the provided facts and any uploaded images. DO NOT invent numbers that are not given.

VERY IMPORTANT:
- Your narratives MUST explicitly reference the specific observations: pre-cement notes, cleanup behavior, bump pressure,
  bleed-off pattern, final stabilized pressure, flowback volume/behavior, post-job tests, and any mismatch ID information.
- Do NOT write generic boilerplate. Each incident must read as specific to the data for that job.
- If flowback is minimal and pressure drops directly to 0 MPa, strongly consider failure_prior_to_cementing, debris_on_collar, or mismatched_receptacle.
- If there is continued inflow and pressure stabilizes at non-zero MPa, strongly consider third_party_integrity and compressibility_ballooning.
- If mismatch IDs are provided, you MUST comment on their relevance in your cause analysis.
- If images are provided (screenshots, charts, photos), you must interpret them and reference any relevant features (e.g. pressure trend, volumes, TOC indications) in your cause analysis and narratives.

ROOT CAUSE MODULES:
You may ONLY choose from this exact list (zero or more):
{allowed_modules}

MEANING OF MODULES (for your reasoning):
- incorrect_pumping_volume:
  Pumped nominal bump volume but plug did not seal or hold, or pattern matches under/over displacement concerns.
- compressibility_ballooning:
  Flowback that may be explained by fluid compressibility, trapped-volume expansion, ballooning, thermal rebound.
- failure_prior_to_cementing:
  Debris / glass / buoyancy-disc fragments / damage that interfered with plug landing BEFORE cementing.
- mismatched_receptacle:
  Plug nose ID and float collar receptacle ID do not match, so plug can land but not seal.
- debris_on_collar:
  Debris sitting on the float collar landing face so plug cannot seat; bridge plug / packer later tests casing above successfully.
- third_party_integrity:
  Leak path above the float equipment (toe subs, sleeves, connections, other accessories), not float equipment failure itself.

COMPRESSIBILITY OUTCOME:
Return either:
- "plausible"       -> if observed flowback/inflow could be normal compressibility / ballooning / thermal rebound.
- "exceeds_normal"  -> if flowback volume and continued inflow clearly indicate a communication path above the float equipment.

NARRATIVE SECTIONS (MAKE THEM JUICY):
Return 4 long-form fields in "narrative_sections":

1) incident_summary:
   - 3 to 6 paragraphs.
   - Walk through the whole sequence chronologically: drilling, running casing/accessories, buoyancy/airlock or toe subs,
     cleanup circulation, cement pumping, displacement, plug drop/bump attempts, flowback/inflow observations.
   - Refer explicitly to the key numbers (depths, volumes, pressures, rates, any meaningful timing).
   - Comment on abnormal SPP, pump behavior, or well response.

2) incident_review:
   - 1 to 3 paragraphs.
   - Focus on QA/QC and process: how the equipment is normally inspected, tested, and built.
   - If relevant, mention debris-handling design features, valve design, latch design, typical testing practices,
     and how they relate to this job.

3) conclusion:
   - 2 to 4 paragraphs.
   - Clearly state the most probable cause(s) of the incident given the data, referencing the selected root cause modules.
   - Explain how the observed pressures, volumes, bump signatures, and flowback support or contradict each possible cause.
   - Provide clear recommendations for follow-up, data review, or procedural changes.

4) overall_cause_analysis:
   - 3 to 6 paragraphs of deeper engineering reasoning.
   - Compare the observed behavior to what you would expect for:
        * a clean, successful bump with good integrity,
        * pure compressibility/thermal effects,
        * leaks across or above the float equipment,
        * debris / damage / mismatched hardware at the latch.
   - Step-by-step logic: what the data says, what it rules out, and why the final selected causes are most consistent.

OUTPUT FORMAT:
Return STRICT JSON with keys:
{{
  "root_cause_blocks": [...],
  "compressibility_outcome": "plausible" | "exceeds_normal",
  "narrative_sections": {{
    "incident_summary": "string",
    "incident_review": "string",
    "conclusion": "string",
    "overall_cause_analysis": "string"
  }}
}}

CRITICAL:
- 'root_cause_blocks' MUST ONLY contain items from this allowed set: {allowed_modules}
- Do NOT include any extra keys.
- No markdown, no headings like "INCIDENT SUMMARY:" inside the strings themselves.
- Just plain text paragraphs in each string, separated by blank lines where natural.
- Do NOT mention these instructions in your output.
""".strip()

    # Build the user content: text + optional images (OpenAI-style content array)
    # Allow user_content to be a list of dicts with potentially nested dicts as values
    user_content: list[dict[str, object]] = [
        {
            "type": "text",
            "text": f"FACTS:\n{facts_blob}\n\nIf images are present below, use them to refine your assessment.\n"
        }
    ]

    for img in images:
        data_url = f"data:{img['mime_type']};base64,{img['b64']}"
        user_content.append({
            "type": "image_url",
            "image_url": {"url": data_url}
        })

    url = "https://api.x.ai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": user_content},
        ],
        "temperature": 0.25,
        "max_tokens": 1800,
    }

    try:
        resp = requests.post(url, json=payload, headers=headers, timeout=90)
        resp.raise_for_status()
        data = resp.json()
        ai_text = data["choices"][0]["message"]["content"].strip()
        parsed = json.loads(ai_text)
        return parsed
    except Exception as e:
        # Fallback if Grok fails — keep report generation alive
        return {
            "root_cause_blocks": ["third_party_integrity", "compressibility_ballooning"],
            "compressibility_outcome": "exceeds_normal",
            "narrative_sections": {
                "incident_summary": f"[AI FAILED, placeholder summary: {e}]",
                "incident_review": (
                    "CCAI inspects, builds, and delivers float equipment and plugs under "
                    "documented QA/QC processes. All assemblies are visually inspected, "
                    "function tested, and labelled prior to delivery."
                ),
                "conclusion": (
                    "Due to an AI generation error, the detailed narrative conclusion could "
                    "not be produced for this revision. Engineering review should focus on "
                    "displacement volumes, bump behavior, bleed-down pattern, and any evidence "
                    "of communication above the float collar."
                ),
                "overall_cause_analysis": (
                    "AI reasoning unavailable for this run. Review Pason EDR, cement charts, "
                    "and post-job tests manually to confirm probable cause."
                )
            }
        }


# ============================================================
# Build report text from user_data + AI result
# ============================================================

def build_report_text(user_data: dict, ai_result: dict):
    # Attach Grok decisions back onto data so our templates can see them
    user_data["compressibility_outcome"] = ai_result.get("compressibility_outcome", "exceeds_normal")

    # 1. Header
    accessories_joined = ", ".join(user_data["accessories"])
    short_desc = user_data["string_desc"]
    if len(short_desc) > 60:
        short_desc = short_desc[:57] + "..."
    header_text = f"""ENGINEERING REPORT

{user_data['date_of_report']}

Revision: {user_data['revision']}
Report: {user_data['cir_number']}

{user_data['customer']} – {short_desc}: {user_data['title_line']}

INCIDENT REPORT NUMBER: {user_data['cir_number']}
CUSTOMER: {user_data['customer']}
AUTHOR: {user_data['author']}
Rig: {user_data['rig']}
SURFACE: {user_data['surface_location']}
UWI: {user_data['uwi']}
EQUIPMENT: {accessories_joined}
""".strip()

    # 2. Incident summary (Grok text)
    incident_summary_text = "INCIDENT SUMMARY\n\n" + ai_result["narrative_sections"]["incident_summary"].strip()

    # 3. Volume / depth summary (from numbers)
    vt = user_data["volume_table"]
    volume_table_text = f"""VOLUME / DEPTH SUMMARY

Well TD: {vt.get('well_td_m', 'N/A')} m
Well TVD: {vt.get('well_tvd_m', 'N/A')} m
Shoe Depth: {vt.get('shoe_depth_m', 'N/A')} m
Float Collar Top Depth: {vt.get('float_collar_depth_m', 'N/A')} m
Airlock / CBS Depth: {vt.get('airlock_depth_m', 'N/A')} m
Crossover Depth: {vt.get('crossover_depth_m', 'N/A')} m

Casing Volume Calculated:
  Nominal: {vt.get('casing_vol_nominal_m3', 'N/A')} m³
  Min:     {vt.get('casing_vol_min_m3', 'N/A')} m³
  Max:     {vt.get('casing_vol_max_m3', 'N/A')} m³

Volume to Airlock / CBS: {vt.get('volume_to_airlock_m3', 'N/A')} m³
Buoyant Volume (nominal): {vt.get('buoyant_volume_nom_m3', 'N/A')} m³

Displacement Volume Pumped: {vt.get('displacement_pumped_m3', 'N/A')} m³
Excess Cement to Surface:  {vt.get('excess_to_surface_m3', 'N/A')} m³
""".strip()

    # 4. Incident review (Grok)
    incident_review_text = "INCIDENT REVIEW\n\n" + ai_result["narrative_sections"]["incident_review"].strip()

    # 5. Potential root causes (our templates, based on Grok selection)
    root_cause_sections = []
    root_block_keys = ai_result.get("root_cause_blocks", [])
    for block_key in root_block_keys:
        builder = ROOT_CAUSE_BLOCK_BUILDERS.get(block_key)
        if builder:
            temp = {
                **user_data,
                "post_job": user_data["post_job"],
                "volume_table": user_data["volume_table"],
                "compressibility_outcome": user_data.get("compressibility_outcome", "exceeds_normal"),
            }
            root_cause_sections.append(builder(temp))

    analysis_extra = ai_result["narrative_sections"].get("overall_cause_analysis", "").strip()
    if analysis_extra:
        root_cause_sections.append(
            "Engineering Assessment / Cause Analysis\n\n" + analysis_extra
        )

    potential_root_causes_text = "POTENTIAL ROOT CAUSES\n\n" + "\n\n".join(root_cause_sections).strip()

    # 6. Conclusion (Grok)
    conclusion_text = "CONCLUSION\n\n" + ai_result["narrative_sections"]["conclusion"].strip()

    parts = [
        header_text,
        "",
        incident_summary_text,
        "",
        volume_table_text,
        "",
        incident_review_text,
        "",
        potential_root_causes_text,
        "",
        conclusion_text,
        ""
    ]

    return "\n".join(parts).strip() + "\n"


# ============================================================
# DOCX STYLING + RENDERING (to BytesIO) + IMAGE APPENDIX
# ============================================================

def ensure_styles(doc):
    styles = doc.styles

    if "TitleStyle" not in [s.name for s in styles]:
        style = styles.add_style("TitleStyle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["TitleStyle"]
    style.font.name = "Calibri"
    style.font.size = Pt(14)
    style.font.bold = True
    style.paragraph_format.space_after = Pt(6)

    if "SectionHeader" not in [s.name for s in styles]:
        style = styles.add_style("SectionHeader", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["SectionHeader"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)
    style.font.bold = True
    style.all_caps = True
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    if "BodyText" not in [s.name for s in styles]:
        style = styles.add_style("BodyText", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["BodyText"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.bold = False
    style.paragraph_format.space_after = Pt(6)

    if "SubHeader" not in [s.name for s in styles]:
        style = styles.add_style("SubHeader", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["SubHeader"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(3)

    if "MonoBlock" not in [s.name for s in styles]:
        style = styles.add_style("MonoBlock", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["MonoBlock"]
    style.font.name = "Consolas"
    style.font.size = Pt(10)
    style.font.bold = False
    style.paragraph_format.left_indent = Pt(18)
    style.paragraph_format.space_before = Pt(3)
    style.paragraph_format.space_after = Pt(3)


def split_report_into_structures(report_text):
    blocks = [b.strip() for b in report_text.split("\n\n") if b.strip()]
    structured = {
        "header": [],
        "incident_summary": [],
        "volume_table": [],
        "incident_review": [],
        "root_causes": [],
        "conclusion": []
    }

    current_section = None
    in_root_causes = False
    pending_cause_title = None
    pending_cause_body = []

    def flush_pending_cause():
        nonlocal pending_cause_title, pending_cause_body
        if pending_cause_title is not None:
            structured["root_causes"].append({
                "title": pending_cause_title,
                "body_lines": pending_cause_body[:],
            })
        pending_cause_title = None
        pending_cause_body = []

    for block in blocks:
        first_line = block.split("\n", 1)[0].strip()

        if first_line.startswith("ENGINEERING REPORT"):
            current_section = "header"
            structured["header"].append(block)

        elif first_line == "INCIDENT SUMMARY":
            current_section = "incident_summary"
            structured["incident_summary"].append(block)

        elif first_line == "VOLUME / DEPTH SUMMARY":
            current_section = "volume_table"
            structured["volume_table"].append(block)

        elif first_line == "INCIDENT REVIEW":
            current_section = "incident_review"
            structured["incident_review"].append(block)

        elif first_line == "POTENTIAL ROOT CAUSES":
            current_section = "root_causes"
            in_root_causes = True

        elif first_line == "CONCLUSION":
            flush_pending_cause()
            current_section = "conclusion"
            in_root_causes = False
            structured["conclusion"].append(block)

        else:
            if in_root_causes:
                flush_pending_cause()
                lines = block.split("\n")
                pending_cause_title = lines[0].strip()
                pending_cause_body = [ln for ln in lines[1:] if ln.strip() != ""]
            else:
                if current_section == "header":
                    structured["header"].append(block)
                elif current_section == "incident_summary":
                    structured["incident_summary"].append(block)
                elif current_section == "volume_table":
                    structured["volume_table"].append(block)
                elif current_section == "incident_review":
                    structured["incident_review"].append(block)
                elif current_section == "conclusion":
                    structured["conclusion"].append(block)

    flush_pending_cause()
    return structured


def build_docx_bytes(report_text, images=None, filename_hint="output_incident_report.docx") -> BytesIO:
    """
    Render the sectioned report text into a styled .docx and return as BytesIO.
    Also optionally append uploaded images as an APPENDIX section.
    """
    if images is None:
        images = []

    doc = Document()
    ensure_styles(doc)
    data = split_report_into_structures(report_text)

    # HEADER
    for block in data["header"]:
        lines = block.split("\n")
        for j, line in enumerate(lines):
            if j == 0:
                p = doc.add_paragraph(style="TitleStyle")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line.strip())
                run.bold = True
            else:
                p = doc.add_paragraph(style="BodyText")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run(line.strip())
    doc.add_paragraph(" ", style="BodyText")

    # INCIDENT SUMMARY
    if data["incident_summary"]:
        doc.add_paragraph("INCIDENT SUMMARY", style="SectionHeader")
        for block in data["incident_summary"]:
            lines = block.split("\n")
            if lines[0].strip().upper() == "INCIDENT SUMMARY":
                lines = lines[1:]
            for bl in lines:
                if bl.strip():
                    p = doc.add_paragraph(style="BodyText")
                    p.add_run(bl.strip())

    # VOLUME / DEPTH SUMMARY
    if data["volume_table"]:
        doc.add_paragraph("VOLUME / DEPTH SUMMARY", style="SectionHeader")
        for block in data["volume_table"]:
            lines = block.split("\n")
            if lines[0].strip().upper() == "VOLUME / DEPTH SUMMARY":
                lines = lines[1:]
            for bl in lines:
                if bl.strip():
                    p = doc.add_paragraph(style="MonoBlock")
                    p.add_run(bl.rstrip())

    # INCIDENT REVIEW
    if data["incident_review"]:
        doc.add_paragraph("INCIDENT REVIEW", style="SectionHeader")
        for block in data["incident_review"]:
            lines = block.split("\n")
            if lines[0].strip().upper() == "INCIDENT REVIEW":
                lines = lines[1:]
            for bl in lines:
                if bl.strip():
                    p = doc.add_paragraph(style="BodyText")
                    p.add_run(bl.strip())

    # POTENTIAL ROOT CAUSES
    if data["root_causes"]:
        doc.add_paragraph("POTENTIAL ROOT CAUSES", style="SectionHeader")
        for cause in data["root_causes"]:
            p = doc.add_paragraph(style="SubHeader")
            p.add_run(cause["title"])
            for bl in cause["body_lines"]:
                if bl.strip():
                    body_p = doc.add_paragraph(style="BodyText")
                    body_p.add_run(bl.strip())

    # CONCLUSION
    if data["conclusion"]:
        doc.add_paragraph("CONCLUSION", style="SectionHeader")
        for block in data["conclusion"]:
            lines = block.split("\n")
            if lines[0].strip().upper() == "CONCLUSION":
                lines = lines[1:]
            for bl in lines:
                text = bl.strip()
                if not text:
                    continue
                if text.upper() == "DRILLOUT DE-BRIEF":
                    p = doc.add_paragraph(style="SubHeader")
                    p.add_run("DRILLOUT DE-BRIEF")
                else:
                    p = doc.add_paragraph(style="BodyText")
                    p.add_run(text)

    # APPENDIX – IMAGES
    if images:
        doc.add_page_break()
        doc.add_paragraph("APPENDIX – JOB IMAGES", style="SectionHeader")
        for idx, img in enumerate(images, start=1):
            caption = img.get("filename", f"Image {idx}")
            # caption
            cap_p = doc.add_paragraph(style="BodyText")
            cap_p.add_run(caption)
            # picture
            raw = base64.b64decode(img["b64"])
            run = doc.add_paragraph().add_run()
            run.add_picture(BytesIO(raw), width=Inches(5))

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def parse_float_or_none(text: str):
    """
    Helper for manual input mode.
    - "" -> None
    - numeric string -> float
    - anything else -> original string (so we don't crash on '18-20')
    """
    if text is None:
        return None
    text = str(text).strip()
    if not text:
        return None
    try:
        return float(text)
    except ValueError:
        return text


def main():
    st.set_page_config(page_title="Incident Report Builder", layout="wide")

    st.title("Incident Report Builder")

    # Sidebar: config
    st.sidebar.header("Grok Configuration")
    api_key = st.sidebar.text_input("GROK_API_KEY", type="password")
    model = st.sidebar.text_input("Model ID", value=DEFAULT_GROK_MODEL)

    st.sidebar.markdown("---")
    mode = st.sidebar.selectbox(
        "Incident data source",
        [
            "Mock Case 1 – Partial bump & inflow",
            "Mock Case 2 – No isolation / debris",
            "Manual entry",
        ],
        index=0,
    )

    st.sidebar.markdown("---")
    st.sidebar.write(
        "Upload any images you want Grok to interpret:\n"
        "• EDR screenshots\n"
        "• Cement charts\n"
        "• Photos from location\n"
        "These will also be embedded as an appendix in the Word report."
    )

    uploaded_files = st.sidebar.file_uploader(
        "Upload images", type=["png", "jpg", "jpeg", "webp"], accept_multiple_files=True
    )

    # ===== INCIDENT INPUT AREA =====
    if mode == "Mock Case 1 – Partial bump & inflow":
        st.subheader("Incident input – Mock Case 1")
        st.markdown("Using hard-coded numbers for a **partial bump with continued inflow** scenario.")
        user_data = get_mock_user_data_case1()

    elif mode == "Mock Case 2 – No isolation / debris":
        st.subheader("Incident input – Mock Case 2")
        st.markdown("Using hard-coded numbers for a **no isolation / suspected debris / mismatch** scenario.")
        user_data = get_mock_user_data_case2()

    else:
        # MANUAL ENTRY MODE
        st.subheader("Incident input – Manual entry")

        st.markdown("#### Header / identity")
        col1, col2, col3 = st.columns(3)
        with col1:
            cir_number = st.text_input("CIR Number", "CIR-25-XX")
            revision = st.text_input("Revision #", "0")
        with col2:
            date_of_report = st.text_input(
                "Date of Report (YYYY-MM-DD)",
                datetime.today().strftime("%Y-%m-%d"),
            )
            author = st.text_input("Author", "Your Name, P.Eng")
        with col3:
            customer = st.text_input("Customer", "")
            rig = st.text_input("Rig", "")

        surface_location = st.text_input("Surface Location", "")
        uwi = st.text_input("UWI", "")
        title_line = st.text_input("Title Line (short incident title)", "Citadel No Bump & Pressure Test")

        st.markdown("#### String description & accessories")
        string_desc = st.text_input(
            "Casing String Description",
            '4-1/2" 22.47kg/m L80 LTC x 5-1/2" 34.23kg/m L80 LTC long string production casing',
        )
        accessories_raw = st.text_area(
            "Accessories (comma-separated)",
            "15K Citadel SV Float Shoe, 15K Citadel SV Float Collar, "
            "Citadel Latch Cement Plug, 2 x NCS Toe Ports, 62 x NCS Frac Sleeves, NCS Airlock",
        )
        accessories = [a.strip() for a in accessories_raw.split(",") if a.strip()]

        st.markdown("#### Pre-cement / well geometry")
        col1, col2, col3 = st.columns(3)
        with col1:
            hole_size_mm = parse_float_or_none(st.text_input("Hole size (mm)", "171"))
        with col2:
            td_mmd = parse_float_or_none(st.text_input("TD (mMD)", "5283.0"))
        with col3:
            set_depth_mmd = parse_float_or_none(st.text_input("Casing set depth / landed depth (mMD)", "5281.19"))

        pre_cement_notes = st.text_area(
            "Pre-cement notes (SPP erratic, inflow/no inflow, debris/glass, etc.)",
            "No report of inflow prior to cement. Standpipe pressure became erratic during cleanup.",
        )
        circulated_volume_m3 = parse_float_or_none(
            st.text_input("Cleanup circulation volume before cement (m³)", "127.61")
        )

        st.markdown("#### Cement job")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            cement_lead_m3 = parse_float_or_none(st.text_input("Cement lead volume (m³)", "82.46"))
        with col2:
            cement_tail_m3 = parse_float_or_none(st.text_input("Cement tail volume (m³) [blank if none]", ""))
        with col3:
            displacement_pumped_m3 = parse_float_or_none(st.text_input("Displacement pumped (m³)", "45.6"))
        with col4:
            pump_rate_m3_per_min = parse_float_or_none(st.text_input("Pump rate (m³/min)", "1.0"))

        st.markdown("#### Pressures / flowback")
        col1, col2, col3 = st.columns(3)
        with col1:
            fcp_mpa = parse_float_or_none(st.text_input("FCP during displacement (MPa)", "18.54"))
        with col2:
            bump_pressure_mpa = parse_float_or_none(st.text_input("Bump pressure reached (MPa)", "22.46"))
        with col3:
            bledoff_to_mpa = parse_float_or_none(st.text_input("Pressure bled down / stabilized at (MPa)", "13.56"))

        flowback_volume = st.text_input(
            "Flowback volume / behavior",
            "≈1.0 m³ and still flowing",
        )

        st.markdown("#### Volume / depth summary")
        col1, col2, col3 = st.columns(3)
        with col1:
            vt_well_td_m = parse_float_or_none(st.text_input("Well TD (m)", "5283"))
            vt_shoe_depth_m = parse_float_or_none(st.text_input("Shoe depth (m)", "5280.72"))
        with col2:
            vt_well_tvd_m = parse_float_or_none(st.text_input("Well TVD (m)", "1624.8"))
            vt_float_collar_depth_m = parse_float_or_none(
                st.text_input("Float collar top depth (m)", "5267.62")
            )
        with col3:
            vt_airlock_depth_m = parse_float_or_none(st.text_input("Airlock / CBS depth (m)", "2558.61"))
            vt_crossover_depth_m = parse_float_or_none(st.text_input("Crossover depth (m)", "1682.08"))

        col1, col2, col3 = st.columns(3)
        with col1:
            vt_nom = parse_float_or_none(st.text_input("Casing vol nominal (m³)", "45.2"))
        with col2:
            vt_min = parse_float_or_none(st.text_input("Casing vol min (m³)", "43.5"))
        with col3:
            vt_max = parse_float_or_none(st.text_input("Casing vol max (m³)", "47.1"))

        col1, col2, col3 = st.columns(3)
        with col1:
            vt_vol_to_airlock = parse_float_or_none(st.text_input("Volume to Airlock/CBS (m³)", "25.1"))
        with col2:
            vt_buoyant_nom = parse_float_or_none(st.text_input("Buoyant volume nominal (m³)", "20.1"))
        with col3:
            vt_excess_surface = parse_float_or_none(
                st.text_input("Excess cement to surface (m³) [blank if n/a]", "")
            )

        st.markdown("#### Post-job / drillout")
        col1, col2, col3 = st.columns(3)
        with col1:
            post_retest_pressure = st.text_input("Retest / rig pump pressure seen (MPa)", "18-20")
        with col2:
            post_bridge_hold_mpa = st.text_input("Bridge plug / packer held MPa", "21")
        with col3:
            post_bridge_hold_min = st.text_input("Bridge plug / packer hold time (min)", "11")

        col1, col2 = st.columns(2)
        with col1:
            post_tag_depth_m = st.text_input("Tag depth after job (mMD)", "N/A")
        with col2:
            post_squeeze_summary = st.text_input("Squeeze summary / remedial volumes", "N/A")

        st.markdown("#### Optional mismatch info (leave blank if not suspected)")
        col1, col2 = st.columns(2)
        with col1:
            mismatch_rec = st.text_input('Float collar receptacle ID (inches)', "")
        with col2:
            mismatch_plug = st.text_input('Plug nose ID (inches)', "")

        # Build user_data dict matching the mock structure
        user_data = {
            "cir_number": cir_number,
            "date_of_report": date_of_report,
            "revision": revision,
            "customer": customer,
            "rig": rig,
            "surface_location": surface_location,
            "uwi": uwi,
            "author": author,
            "title_line": title_line,
            "string_desc": string_desc,
            "accessories": accessories,
            "hole_size_mm": hole_size_mm,
            "td_mmd": td_mmd,
            "set_depth_mmd": set_depth_mmd,
            "pre_cement_notes": pre_cement_notes,
            "circulated_volume_m3": circulated_volume_m3,
            "cement_lead_m3": cement_lead_m3,
            "cement_tail_m3": cement_tail_m3,
            "displacement_pumped_m3": displacement_pumped_m3,
            "pump_rate_m3_per_min": pump_rate_m3_per_min,
            "fcp_mpa": fcp_mpa,
            "bump_pressure_mpa": bump_pressure_mpa,
            "bledoff_to_mpa": bledoff_to_mpa,
            "flowback_volume": flowback_volume,
            "volume_table": {
                "well_td_m": vt_well_td_m,
                "well_tvd_m": vt_well_tvd_m,
                "shoe_depth_m": vt_shoe_depth_m,
                "float_collar_depth_m": vt_float_collar_depth_m,
                "airlock_depth_m": vt_airlock_depth_m,
                "crossover_depth_m": vt_crossover_depth_m,
                "casing_vol_nominal_m3": vt_nom,
                "casing_vol_min_m3": vt_min,
                "casing_vol_max_m3": vt_max,
                "volume_to_airlock_m3": vt_vol_to_airlock,
                "buoyant_volume_nom_m3": vt_buoyant_nom,
                "displacement_pumped_m3": displacement_pumped_m3,
                "excess_to_surface_m3": vt_excess_surface,
            },
            "post_job": {
                "retest_pressure_mpa": post_retest_pressure,
                "bridge_plug_hold_mpa": post_bridge_hold_mpa,
                "bridge_plug_hold_min": post_bridge_hold_min,
                "tag_depth_m": post_tag_depth_m,
                "squeeze_summary": post_squeeze_summary,
            },
            "mismatch_receptacle_receptacle_id_in": mismatch_rec,
            "mismatch_receptacle_plug_nose_id_in": mismatch_plug,
        }

    # ===== SNAPSHOT + GENERATION =====
    st.subheader("Incident Snapshot")
    st.json(user_data)

    st.markdown("When you're happy, click **Generate Report** to call Grok and build the Word file.")

    generate_button = st.button("Generate Report")

    if generate_button:
        if not api_key:
            st.error("Please enter your GROK_API_KEY in the sidebar.")
            return

        with st.spinner("Calling Grok and generating report..."):
            images_payload = encode_uploaded_images(uploaded_files) if uploaded_files else []
            ai_result = generate_ai_full_report(
                user_data, api_key=api_key, model=model, images=images_payload
            )
            report_text = build_report_text(user_data, ai_result)
            docx_bytes = build_docx_bytes(report_text, images=images_payload)

        st.success("Report generated.")

        with st.expander("Show raw AI JSON (root causes & narratives)", expanded=False):
            st.json(ai_result)

        with st.expander("Show assembled report text (for debugging)", expanded=False):
            st.text(report_text)

        st.download_button(
            label="Download Word Report (.docx)",
            data=docx_bytes,
            file_name=f"{user_data['cir_number'].replace(' ', '_')}_incident_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

if __name__ == "__main__":
    main()
